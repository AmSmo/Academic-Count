import docx
import re
import pprint
wordcount = r"\S+"

this_file = 'article3.docx'
#this_file = 'article3.docx'
#this_file= 'student.docx'
#end_word = False
highlighted_copy = docx.Document()

#def stop_word():
#    answer = input("Would you like to stop at a certain section? Yes or No: ")
#    if answer.lower() == "yes" or answer.lower() == "y":
#        end_word = True
#        return input("What word would you like to stop at? (Note: Case and punctuation sensitive) ")
#    else:
#        end_word = False
class AnalyzeDoc:


    def __init__(self, file, last_word):
        self.no_doubles = ""
        self.last_word=last_word
        self.file = file
        self.opened = self.open_file()
        self.full_text = self.textify()
        self.bibstart = self.get_end_citations()[2]
        self.bibend = self.get_end_citations()[3]
        self.abstract_count = self.get_abstract()[0]
        self.abstract_text = self.get_abstract()[1]
        self.keywords_count = self.get_keyword()[0]
        self.keywords_text = self.get_keyword()[1]
        self.total = self.all_words()
        self.table_count = self.get_tables()[0]
        self.table_text = self.get_tables()[1]
        self.table_intro_count = self.get_table_intros()[0]
        self.table_intro_text = self.get_table_intros()[1]

        self.bibliography_count = self.get_end_citations()[0]
        self.bibliography_text = self.get_end_citations()[1]
        self.figures_count = self.get_figures()[0]
        self.figures_text = self.get_figures()[1]
        self.notes_count = self.get_notes()[0]
        self.notes_text = self.get_notes()[1]
        self.appendices_count = self.get_appendix()[0]
        self.appendices_text = self.get_appendix()[1]
        self.title_count = self.no_title_page()[0]
        self.title_text = self.no_title_page()[1]
        self.citations_count = self.get_in_paper_citations()[0]
        self.citations_text = self.get_in_paper_citations()[1]
        self.false_negatives = self.get_in_paper_citations()[2]
        self.citation_excise_text = self.get_in_paper_citations()[3]
        self.edited=self.full_text
        #self.highlighted = self.highlights()




    def open_file(self):
        with open(self.file, 'rb') as file:
            return docx.Document(file)

    def textify(self):
        entire_doc = []

        for paragraph in self.opened.paragraphs:
            if self.last_word:
                check_end = re.compile(rf'{self.last_word}')
                if re.match(check_end, paragraph.text):
                    break
                else:
                    entire_doc.append(paragraph.text)
            else:
                entire_doc.append(paragraph.text)
        self.paragraphs = entire_doc
        return "\n".join(entire_doc)

    def get_abstract(self):
        abstract_total = 0
        abstract_text = []
        regex_abstract = re.compile(r'(:?Abstract|ABSTRACT)[ :]? ?\n.*.*')
        abstract = re.search(regex_abstract, self.full_text)
        nextsection = r'(\n\n|\n\S+\n|\n\w+:|\n\t\w+:)'

        try:
            post_abs = re.search(nextsection, self.full_text[abstract.start() +12:])

            abstract_text.append(self.full_text[abstract.start():abstract.start() + (post_abs.start()+12)])
            for item in abstract_text:
                abstract_total += len(re.findall(wordcount, item))
            return abstract_total, abstract_text
        except:
            abstract_text.append("<I>Nothing was found as your Abstract</I>")
            return abstract_total, abstract_text

    def no_title_page(self):
        title_text = []
        title_total = 0
        regex_abstract = re.compile(r'(:?Abstract|ABSTRACT)[ :]? ?\n.*.*')
        abstract = re.search(regex_abstract, self.full_text)
        if abstract == None:
            regex_title = re.compile(r'\w+\n\n\n')
            abstract = re.search(regex_title, self.full_text)
        try:

            title_text.append(self.full_text[:abstract.start()])
            for item in title_text:
                title_total+=len(re.findall(wordcount,item))
            return title_total, title_text
        except:
            title_text.append("<I>Nothing was found as your Title Page</I>")
            return title_total, title_text

    def get_keyword(self):
        keywords_total = 0
        keywords_text = []
        regex_keyword = re.compile(r'(Keywords?|KEYWORDS?)[: ]?\n?.*\n')
        keywords = re.search(regex_keyword, self.full_text)
        nextsection = r'(\n\n|\n\w+[ :\n][^\w])'

        try:
            post_key = re.search(nextsection, self.full_text[keywords.start() +12:])

            keywords_text.append(self.full_text[keywords.start():keywords.start() + (post_key.start()+12)])
            for item in keywords_text:
                keywords_total += len(re.findall(wordcount, item))
            return keywords_total, keywords_text
        except:
            keywords_text.append("<I>Nothing was found as Keywords</I>")
            return keywords_total, keywords_text

        # for keyword in keywords:
        #     keywords_total += len(re.findall(wordcount, keyword))
        #     keywords_text.append(keyword)
        # return keywords_total, keywords_text

    def get_figures(self):
        figures_total = 0
        figures_text = []
        regex_figure = r'\n+\t?Figure \d{1,3}\.?.*\n'
        figures = re.findall(regex_figure, self.full_text)

        for figure in figures:
            figures_total += len(re.findall(wordcount,figure))
            figures_text.append(figure)
        return figures_total, figures_text

    def get_table_intros(self):
        tables_intro_total = 0
        tables_intro_text = []
        regex_table = r'\nTable \d\.?.*\n.*'
        tables = re.findall(regex_table, self.full_text)
        for table in tables:
            tables_intro_total += len(re.findall(wordcount,table))
            tables_intro_text.append(table)
        return tables_intro_total, tables_intro_text

    def get_end_citations(self):
        bib_count = 0
        full_bib = ""
        biblography = r'\n(References:?|Bibliography:?)'
        bibstart = re.search(biblography, self.full_text)
        nextsection= r'\n\n|\Z'


        try:
            pre_bib = bibstart.start()
            post_bib = re.search(nextsection, self.full_text[bibstart.start()+10:])

            full_bib+=(self.full_text[bibstart.start():bibstart.start()+post_bib.start()+10])
            bib_count += len(re.findall(wordcount, full_bib))
            bibend = bibstart.start()+post_bib.start()+10

            return bib_count, full_bib, pre_bib, bibend
        except:
            pre_bib = len(self.full_text)
            bibend = len(self.full_text)
            return bib_count, full_bib, pre_bib, bibend

    def get_in_paper_citations(self):
        citations_total = 0
        citations_text = []
        false_negatives = []
        excised_word = []
        regex_citations = r'\(.*?\)'
        citations_not_bibliography=(self.full_text[:self.bibstart]+(self.full_text[self.bibend:]))
        citations = re.findall(regex_citations, citations_not_bibliography)
        for citation in citations:
            if re.search(r'(?<!-)[12][0-9]\d{2}(?! ms|/| milliseconds)', citation):
                multicaps = re.search(r'[A-Za-z .]*?[A-Z]{2,}[A-Za-z]*\s?\w+[:,;]?', citation)
                doublecap = re.search(r'[A-Z]{2}', citation)
                sentence = re.search(r'[\( ][a-z ,]+', citation)
                examples = re.search(r'\((i.e.|e.g.).*[,;]', citation)
                if multicaps:
                    excised_word.append(f"{multicaps.group()} FROM: {citation}")
                    citations_total -= len(re.findall(wordcount, multicaps.group()))
                    citation = citation.replace(multicaps.group(), "")
                elif examples:
                    excised_word.append(f"{examples.group()} FROM: {citation}")

                    citations_total -= len(re.findall(wordcount, examples.group()))
                    citation = citation.replace(examples.group(), "")
                elif doublecap:
                    excised_word.append(f"{doublecap.group()} FROM: {citation}")

                    citations_total -= len(re.findall(wordcount, doublecap.group()))
                    citation = citation.replace(doublecap.group(), "")
                elif sentence and sentence.group() != " et al" and sentence.group() != " de "\
                        and sentence.group() != " van de ":
                    excised_word.append(f"{sentence.group()} FROM: {citation}")

                    citations_total -= len(re.findall(wordcount, sentence.group()))
                    citation = citation.replace(sentence.group(), "")
                citations_total += len(re.findall(wordcount, citation))
                citations_text.append(citation)
            else:
                false_negatives.append(citation)
        return citations_total, citations_text, false_negatives, excised_word




    def all_words(self):
        return len(re.findall(wordcount, self.full_text))

    def get_tables(self):
        tableread=[]
        for table in self.opened.tables:
            current=""
            prohibit_splitcells=[]
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                     if cell.text not in prohibit_splitcells[-4:]:
                        current += cell.text + " "
                        prohibit_splitcells.append(cell.text)
            tableread.append(current)
        tablecount= r"[\(\)a-zA-z0-9.-/æ=]+"
        table_total=0
        for table in tableread:
            table_total+=(len(re.findall(tablecount, table)))
        return table_total, tableread

    def get_notes(self):
        notes_total = 0
        notes_text = []
        regex_notes = re.compile(r'\nNotes?:?')
        notes_start = re.search(regex_notes, self.full_text)
        nextsection = r'\n\n|\Z'

        try:
            post_notes = re.search(nextsection, self.full_text[notes_start.start():])
            notes_text.append(self.full_text[notes_start.start():notes_start.start() + post_notes.end()])
            next_search = self.full_text[notes_start.start() + post_notes.end():]
            next_note = re.compile(r"\nNote")
            while re.findall(next_note, next_search):
                next_start = (re.search(next_note, next_search))
                post_notes = re.search(nextsection, next_search[next_start.end():])

                notes_text.append(next_search[next_start.start():next_start.start()+post_notes.end()+5])
                next_search = next_search[next_start.start() + post_notes.end():]

            for note in notes_text:
                notes_total += len(re.findall(wordcount, note))
            return notes_total, notes_text
        except:

            return notes_total, notes_text

    def get_appendix(self):
        appendix_total = 0
        appendix_text = []
        regex_appendix = re.compile(r'(Appendix|Appendices)\n')
        appendix_start = re.search(regex_appendix, self.full_text)
        nextsection = r'\n\n|\Z'
        try:
            post_appendix = re.search(nextsection, self.full_text[appendix_start.end()+(len(appendix_start.group())):])
            print
            appendix_text.append(self.full_text[appendix_start.start():appendix_start.start() + post_appendix.end()+(len(appendix_start.group()))])
            next_search = self.full_text[appendix_start.start() + post_appendix.end()+(len(appendix_start.group())):]
            next_appendix = re.compile(r"\nAppen")
            while re.findall(next_appendix, next_search):
                next_start=(re.search(next_appendix,next_search))

                post_appendix=re.search(nextsection,next_search[next_start.end():])
                appendix_text.append(next_search[next_start.start():post_appendix.end()+(len(appendix_start.group()))])
                next_search= next_search[next_start.start()+post_appendix.end():]
            for app in appendix_text:
                appendix_total += len(re.findall(wordcount, app))

            return appendix_total, appendix_text
        except:
            for app in appendix_text:
                appendix_total += len(re.findall(wordcount, app))
            return appendix_total, appendix_text

    def highlights(self):
         for text in self.abstract_text:
            if text in self.edited:
                  self.edited = self.edited.replace(text, f"<mark>{text}</mark>")
                  self.edited = self.edited.replace("\n", "<br><br>")
         return self.edited


    # for table in read.tables:
    #     df = [['' for i in range(len(table.columns))] for j in range(len(table.rows))]
    #     for i, row in enumerate(table.rows):
    #         for j, cell in enumerate(row.cells):
    #             if cell.text:
    #                 df[i][j] = cell.text
    #     tables.append(pd.DataFrame(df))
    #     table.columns.name = df.iloc[0]
    #     table.columns.name=df.iloc[0]


# test= AnalyzeDoc(this_file, None)
#
# print(test.citation_excise_text)
